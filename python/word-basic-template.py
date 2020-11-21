from docx import Document
from docx.shared import Inches

document = Document()

"""在文档末尾添加标题行.

第一个参数是 *文本* 
第二个参数是大纲基本 *基本*. 
级别为0是风格设置为`Title`. 如果级别为1或空，使用`Heading 1`，否则为指定级别.
"""
document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
"""
第一个参数是 *文本* 
第二个参数是风格 *基本*.
特殊字符
tab (``\\t``) 制表位
newline (``\\n``) 新行
carriage (``\\r``) 断行
"""
p.add_run('bold').bold = True # 属性设置函数，定义在text/run.py
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)


"""
段落的样式存在anaconda3\lib\site-packages\docx\templates\default-styles.xml
第一个参数是 *文本* 
第二个参数是风格 *基本*.
特殊字符
tab (``\\t``) 制表位
newline (``\\n``) 新行
carriage (``\\r``) 断行
"""
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('shot1.png', width=Inches(1.25))

# 添加表格数据
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
# 标题行数据
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
# 遍历数据，诸葛插入
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

# 添加分页符
document.add_page_break()

document.save('demo.docx')