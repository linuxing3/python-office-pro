# 坑：install python-docx,而不是docx
from docx import Document
# 坑：WD_ALIGN_PARAGRAPH是常量，但类名是另一个
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches # 磅数和英寸
from docx.oxml.ns import qn # 中文格式

company_list = ['one', 'two']

for i in company_list:
  # 新建文档
  document = Document()
  document.styles['Normal'].font.name = u'宋体'
  # 要修改默认的中文和西文字体
  document.styles['Normal'].element.rPr.rFonts.set(qn('w:estaAsia'), u'宋体')
  # 添加段落
  p1 = document.add_paragraph()
  # 左对齐
  p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
  # 段落内输入文字
  run1 = p1.add_run('Some text')
  # 设置字体
  run1.font.name = '微软雅黑'
  run1.font.size = Pt(21)
  run1.font.bold = True
  run1.space_after = Pt(5)
  run1.space_before = Pt(5)
  document.add_page_break()
  # write Document
  document.save('%s.docx' % i)